import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 引入我们的路径工具，解决打包后的模板寻址问题
from src.utils.path_utils import get_resource_path


def add_dynamic_page_numbers(paragraph, font_name):
    """为 Word 页脚添加动态页码 (第 X 页 / 共 Y 页)"""

    def add_run(text=""):
        r = paragraph.add_run(text)
        r.font.name = font_name
        getattr(r, "_element").rPr.rFonts.set(qn('w:eastAsia'), font_name)
        r.font.color.rgb = RGBColor(128, 128, 128)
        return r

    def add_field(r, code):
        f1 = OxmlElement('w:fldChar')
        f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = code
        f2 = OxmlElement('w:fldChar')
        f2.set(qn('w:fldCharType'), 'end')
        getattr(r, "_r").append(f1)
        getattr(r, "_r").append(it)
        getattr(r, "_r").append(f2)

    add_run("第 ")
    add_field(add_run(), "PAGE")
    add_run(" 页 / 共 ")
    add_field(add_run(), "NUMPAGES")
    add_run(" 页")


def save_output(extracted_data: list, out_folder: str, filename: str, fmt: str,
                t_font: str, t_size: int, b_font: str, b_size: int,
                t_rgb: tuple, b_rgb: tuple, enable_hf: bool = False,
                progress_cb=None) -> str:
    """
    将提取到的代码数据保存为 Word 或 TXT 文件。
    """
    os.makedirs(out_folder, exist_ok=True)
    ext = 'docx' if fmt == 'Word' else 'txt'
    full_path = os.path.join(out_folder, f"{filename}.{ext}")

    if fmt == "Word":
        # 获取基础模板路径（注意这里的相对路径已经指向了新的 assets 目录）
        tp = get_resource_path('assets/templates/quema_template.docx')

        # 判断底包是否存在，若存在且有效则加载，否则用默认空白档
        if os.path.exists(tp) and os.path.getsize(tp) > 100:
            doc = Document(tp)
        else:
            doc = Document()

        # 注入版权与元数据摘要
        doc.core_properties.author = "QueMa_Office"
        doc.core_properties.comments = "源码一键提取工具生成"
        doc.core_properties.title = "源码整理交付物"

        # 设置正文全局默认样式
        style = doc.styles['Normal']
        style.font.name = b_font
        getattr(style, "_element").rPr.rFonts.set(qn('w:eastAsia'), b_font)
        style.font.size = Pt(b_size)

        if enable_hf:
            sec = doc.sections[0]
            # 顶部页眉：注入文件名
            header = sec.header
            h_p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h_r = h_p.add_run(filename)
            h_r.font.name = b_font
            getattr(h_r, "_element").rPr.rFonts.set(qn('w:eastAsia'), b_font)
            h_r.font.color.rgb = RGBColor(128, 128, 128)

            # 底部页脚：动态页码
            f_para = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
            f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_dynamic_page_numbers(f_para, b_font)

        total = len(extracted_data)
        for i, item in enumerate(extracted_data):
            h = doc.add_heading(f"文件：{item['filename']}", level=2)
            for r in h.runs:
                r.font.name = t_font
                getattr(r, "_element").rPr.rFonts.set(qn('w:eastAsia'), t_font)
                r.font.size = Pt(t_size)
                r.font.color.rgb = RGBColor(*t_rgb)

            p = doc.add_paragraph()
            run = p.add_run(item['content'])
            run.font.name = b_font
            getattr(run, "_element").rPr.rFonts.set(qn('w:eastAsia'), b_font)
            run.font.size = Pt(b_size)
            run.font.color.rgb = RGBColor(*b_rgb)

            # 更新排版进度
            if progress_cb:
                progress_cb(40 + int((i / total) * 55))
        doc.save(full_path)

    else:
        # TXT 纯文本模式
        with open(full_path, 'w', encoding='utf-8') as f:
            total = len(extracted_data)
            for i, item in enumerate(extracted_data):
                f.write(f"文件: {item['filename']}\n{'-' * 50}\n{item['content']}\n\n")
                if progress_cb:
                    progress_cb(40 + int((i / total) * 55))

    return full_path


def save_tree_output(tree_str: str, out_folder: str, filename: str, fmt: str) -> str:
    """将目录树专供导出为 Word 或 TXT"""
    os.makedirs(out_folder, exist_ok=True)
    ext = 'docx' if fmt == 'Word' else 'txt'
    full_path = os.path.join(out_folder, f"{filename}.{ext}")

    if fmt == "Word":
        doc = Document()
        doc.core_properties.author = "QueMa_Office"
        doc.core_properties.title = "项目目录结构树"

        # 写入标题
        doc.add_heading("项目目录结构树", level=1)

        # 写入等宽树状图
        para = doc.add_paragraph()
        run = para.add_run(tree_str)
        # 强制使用等宽字体确保分支符号对齐
        run.font.name = 'Consolas'
        getattr(run, "_element").rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(10)

        doc.save(full_path)
    else:
        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(f"【项目目录结构树】\n{'-' * 30}\n")
            f.write(tree_str)

    return full_path